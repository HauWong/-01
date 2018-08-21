# !/usr/bin/env python
# -*- coding:utf-8 -*-

import os
import sys
import time
import docx
import shutil
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler


class EventHandler(FileSystemEventHandler):
    def __init__(self):
        FileSystemEventHandler.__init__(self)
        self.file_manager = FileManager()

    def on_moved(self, event):
        if event.is_directory:
            # print('directory moved from {0} to {1}'.format(event.src_path, event.dest_path))
            pass
        else:
            print('file moved from {0} to {1}'.format(event.src_path, event.dest_path))
            self.file_manager.move_file(event.dest_path)

    def on_created(self, event):
        if event.is_directory:
            # print('directory created: {0}'.format(event.src_path))
            pass
        else:
            print('file created:{0}'.format(event.src_path))
            self.file_manager.add_folder(event.src_path)


class FileManager(object):
    def __init__(self):
        self.path = '.'
        self.file = 'none'

    def add_folder(self, str):
        self.path, self.file = os.path.split(str)
        if os.path.isdir(str):
            return
        file_name, file_ext = self.file.split('.')
        dir_ls = os.listdir(path)
        if file_ext != 'pdf':
            return
        elif file_name in dir_ls:
            return
        elif file_name == path.split('\\')[-1]:
            return
        else:
            folder_name = os.path.join(path, file_name)
            os.mkdir(folder_name)
            shutil.move(str, os.path.join(folder_name, self.file))
            with open(os.path.join(folder_name, 'basic_info.txt'), 'w') as txt_f:
                info = 'name: '+file_name
                txt_f.write(info)

            doc_file = docx.Document()
            doc_file.add_paragraph(file_name)
            doc_file.save(os.path.join(folder_name, 'note.docx'))

    def move_file(self, str):
        self.path, self.file = os.path.split(str)
        if self.file.split('.')[1] == 'pdf':
            dir_ls = os.listdir(self.path)
            for file in dir_ls:
                if file == 'basic_info.txt':
                    with open(os.path.join(self.path, file), 'w') as txt_f:
                        info = 'name: ' + self.file.split('.')[0]
                        txt_f.write(info)
                if file == 'note.docx':
                    doc_file = docx.Document(os.path.join(self.path, file))
                    doc_file.paragraphs[0].text = self.file.split('.')[0]
                    doc_file.save(os.path.join(self.path, file))

            os.rename(self.path, self.file.split('.')[0])
            pass


if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else '.'
    event_handler = EventHandler()
    observer = Observer()
    observer.schedule(event_handler, path, recursive=True)
    observer.start()
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()